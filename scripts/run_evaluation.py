#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Evaluation script for LegalAI pipeline.

Input: contracts_labels.json (rule inventory) + dataset (train_all.jsonl) for ground truth.
Runs: OCR -> normalize -> clauses -> rules + ML on the 22 labeled contracts.
Output: accuracy, precision, recall, F1 (micro + macro), JSON + pretty table.

Run (from project root):
    python scripts/run_evaluation.py           # evaluate on all contracts
    python scripts/run_evaluation.py --strict  # evaluate ONLY on held-out test contracts

Inside Docker:
    docker exec -it <container> python scripts/run_evaluation.py
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from collections import defaultdict

# Add project root to path
ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

import numpy as np
from sklearn.metrics import precision_score, recall_score, f1_score


def _extract_text(path: Path) -> str:
    """Extract text from PDF or DOCX."""
    if not path.exists():
        return ""
    suf = path.suffix.lower()
    if suf == ".pdf":
        if not fitz:
            return ""
        try:
            doc = fitz.open(path)
            parts = []
            for page in doc:
                t = page.get_text("text") or ""
                if t.strip():
                    parts.append(t)
            doc.close()
            return "\n".join(parts).strip()
        except Exception:
            return ""
    if suf == ".docx":
        try:
            from docx import Document
            doc = Document(path)
            parts = [p.text for p in doc.paragraphs if p.text]
            for t in doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        if cell.text:
                            parts.append(cell.text)
            return "\n".join(parts).strip()
        except Exception:
            return ""
    return ""


def _load_ground_truth(dataset_paths: list[Path]) -> tuple[dict[str, set[str]], dict[str, str]]:
    """
    Aggregate labels by contract_id from dataset JSONL files.
    Returns (gt: contract_id -> labels, file_map: contract_id -> filename).
    """
    gt: dict[str, set[str]] = defaultdict(set)
    file_map: dict[str, str] = {}
    for dp in dataset_paths:
        if not dp.exists():
            continue
        with open(dp, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    r = json.loads(line)
                except Exception:
                    continue
                cid = r.get("contract_id") or (r.get("file_name", "") or "").replace(".pdf", "").replace(".docx", "").strip()
                fname = (r.get("file_name") or "").strip()
                if not cid:
                    continue
                if fname.endswith(".pdf"):
                    file_map[cid] = fname
                elif fname.endswith(".docx"):
                    file_map[cid] = fname
                elif cid not in file_map:
                    file_map[cid] = f"{cid}.pdf"
                labels = r.get("labels") or []
                for lab in labels:
                    if lab:
                        gt[cid].add(str(lab).strip())
    return dict(gt), file_map


def _coerce_hits(raw) -> list:
    if raw is None:
        return []
    if isinstance(raw, list):
        return [x if isinstance(x, dict) else {"description": str(x)} for x in raw]
    if isinstance(raw, dict):
        if "matches" in raw and isinstance(raw["matches"], list):
            return [h if isinstance(h, dict) else {"description": str(h)} for h in raw["matches"]]
        if "hits" in raw:
            h = raw["hits"]
            if isinstance(h, list):
                return [x if isinstance(x, dict) else {"description": str(x)} for x in h]
    return [{"description": str(raw)}]


def _run_pipeline(text: str, rules_dir: Path, laws_dir: Path) -> tuple[list[str], list[dict], bool]:
    """Run rules + ML on text. Returns (rule_ids, ml_predictions, ml_used)."""
    from app.utils_text import normalize_for_rules
    from app.rules import RuleEngine

    text_norm = normalize_for_rules(text)
    rule_ids: list[str] = []
    ml_preds: list[dict] = []
    ml_used = False

    rule_engine = RuleEngine(rules_dir, laws_dir)

    # Rule engine
    try:
        hits = rule_engine.check_text(text_norm, law_scope=["labor", "cross_border"]) or []
        hits = _coerce_hits(hits)
        for h in hits:
            rid = h.get("rule_id") or h.get("id")
            if rid:
                rule_ids.append(str(rid))
    except Exception:
        pass

    # ML predictor
    try:
        from app.ml_predictor import _HAS_ML_PREDICTOR, predict_rule_scores_full
        if _HAS_ML_PREDICTOR and predict_rule_scores_full:
            preds = predict_rule_scores_full(text_norm, sort=True)
            ml_preds = preds
            ml_used = True
            for p in preds:
                if p.get("passed_threshold") and p.get("rule_id"):
                    rule_ids.append(str(p["rule_id"]))
    except Exception:
        pass

    rule_ids = list(dict.fromkeys(rule_ids))
    return rule_ids, ml_preds, ml_used


def _load_held_out_test_ids(base: Path) -> set:
    """Load held-out test contract IDs for strict evaluation."""
    path = base / "data" / "held_out_test_contracts.txt"
    out = set()
    if path.exists():
        for line in path.read_text(encoding="utf-8").splitlines():
            s = line.strip().split("#")[0].strip()
            if s:
                out.add(s)
    return out


def main():
    parser = argparse.ArgumentParser(description="Evaluate LegalAI pipeline")
    parser.add_argument("--strict", action="store_true", help="Evaluate ONLY on held-out test contracts")
    args = parser.parse_args()

    base = ROOT
    contracts_pdf = base / "data" / "contracts_raw" / "pdf"
    contracts_docx = base / "data" / "contracts_raw" / "docx"
    dataset_paths = [
        base / "data" / "dataset" / "train_all.jsonl",
        base / "data" / "dataset" / "silver_dataset.jsonl",
    ]
    labels_file = base / "data" / "contracts_labels.json"

    gt, file_map = _load_ground_truth(dataset_paths)

    if args.strict:
        held_out = _load_held_out_test_ids(base)
        if held_out:
            gt = {k: v for k, v in gt.items() if k in held_out}
            print(f"[STRICT] Evaluating only on {len(gt)} held-out test contracts: {sorted(gt.keys())}")
        else:
            print("[WARN] --strict used but no held_out_test_contracts.txt found; evaluating on all.")
    if not gt:
        print("No ground truth found. Ensure data/dataset/*.jsonl exist.")
        return 1

    rules_dir = base / "rules"
    laws_dir = base / "laws"

    all_rule_ids = set()
    for s in gt.values():
        all_rule_ids.update(s)
    all_rule_ids = sorted(all_rule_ids)

    results = []
    y_true_all = []
    y_pred_all = []

    for contract_id, gt_labels in sorted(gt.items()):
        fn = file_map.get(contract_id) or f"{contract_id}.pdf"
        path = contracts_pdf / fn
        if not path.exists():
            path = contracts_docx / fn.replace(".pdf", ".docx")
        if not path.exists():
            path = contracts_docx / fn
        if not path.exists():
            path = contracts_pdf / contract_id
        if not path.exists():
            print(f"  Skip {contract_id}: file not found ({fn})")
            continue

        text = _extract_text(path)
        if len(text) < 20:
            print(f"  Skip {contract_id}: empty/short text")
            continue

        pred_ids, ml_preds, ml_used = _run_pipeline(text, rules_dir, laws_dir)

        gt_set = set(gt_labels)
        pred_set = set(pred_ids)

        # Binarize for sklearn (per rule_id)
        yt = [1 if r in gt_set else 0 for r in all_rule_ids]
        yp = [1 if r in pred_set else 0 for r in all_rule_ids]
        y_true_all.append(yt)
        y_pred_all.append(yp)

        tp = len(gt_set & pred_set)
        fp = len(pred_set - gt_set)
        fn_count = len(gt_set - pred_set)
        acc = (tp + (len(all_rule_ids) - tp - fp - fn_count)) / max(1, len(all_rule_ids))

        results.append({
            "contract_id": contract_id,
            "gt_count": len(gt_set),
            "pred_count": len(pred_set),
            "tp": tp,
            "fp": fp,
            "fn": fn_count,
            "accuracy": round(acc, 4),
            "ml_used": ml_used,
        })

    if not results:
        print("No contracts evaluated.")
        return 1

    y_true_arr = np.array(y_true_all)
    y_pred_arr = np.array(y_pred_all)

    # Flatten for micro
    yt_flat = y_true_arr.flatten()
    yp_flat = y_pred_arr.flatten()

    metrics = {
        "n_contracts": len(results),
        "accuracy_micro": float(np.mean((yt_flat == yp_flat).astype(float))),
        "precision_micro": float(precision_score(yt_flat, yp_flat, zero_division=0)),
        "recall_micro": float(recall_score(yt_flat, yp_flat, zero_division=0)),
        "f1_micro": float(f1_score(yt_flat, yp_flat, zero_division=0)),
        "precision_macro": float(precision_score(yt_flat, yp_flat, average="macro", zero_division=0)),
        "recall_macro": float(recall_score(yt_flat, yp_flat, average="macro", zero_division=0)),
        "f1_macro": float(f1_score(yt_flat, yp_flat, average="macro", zero_division=0)),
    }

    out = {
        "metrics": metrics,
        "per_contract": results,
    }

    out_path = base / "reports" / "evaluation_results.json"
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(out, f, ensure_ascii=False, indent=2)

    # Pretty table
    mode = " (STRICT - held-out test only)" if args.strict else ""
    print("\n" + "=" * 60)
    print(f"LEGALAI EVALUATION RESULTS{mode}")
    print("=" * 60)
    print(f"\nContracts evaluated: {metrics['n_contracts']}")
    print(f"\nMicro metrics:")
    print(f"  Accuracy:  {metrics['accuracy_micro']:.4f}")
    print(f"  Precision: {metrics['precision_micro']:.4f}")
    print(f"  Recall:    {metrics['recall_micro']:.4f}")
    print(f"  F1:        {metrics['f1_micro']:.4f}")
    print(f"\nMacro metrics:")
    print(f"  Precision: {metrics['precision_macro']:.4f}")
    print(f"  Recall:    {metrics['recall_macro']:.4f}")
    print(f"  F1:        {metrics['f1_macro']:.4f}")
    print(f"\nPer-contract (first 10):")
    for r in results[:10]:
        print(f"  {r['contract_id']}: TP={r['tp']} FP={r['fp']} FN={r['fn']} acc={r['accuracy']:.2f} ml={r['ml_used']}")
    print(f"\nFull results: {out_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
