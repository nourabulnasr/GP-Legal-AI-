from __future__ import annotations

import json
import os
from pathlib import Path
_env_path = Path(__file__).resolve().parent.parent / ".env"
_cwd_env = Path(os.getcwd()) / ".env"
try:
    from dotenv import load_dotenv
    if _env_path.exists():
        load_dotenv(_env_path)
    if _cwd_env.exists() and str(_cwd_env) != str(_env_path):
        load_dotenv(_cwd_env, override=False)
    if not _env_path.exists() and _cwd_env.exists():
        load_dotenv(_cwd_env)
except Exception:
    pass
# Reduce noisy transformer LOAD REPORT / UNEXPECTED keys at startup (expected for AraBERT/MiniLM)
if "TRANSFORMERS_VERBOSITY" not in os.environ:
    os.environ["TRANSFORMERS_VERBOSITY"] = "error"

from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from typing import Any, List, Dict, Optional, Generator
import io
import json
import hashlib
import unicodedata

def _normalize_contract_text(text: str) -> str:
    text = unicodedata.normalize("NFKC", text or "")
    # ... باقي النورمالايز عندك زي ما هو (تشيل تشكيل/تطويل/مسافات..)
    return text

import fitz  # PyMuPDF
from PIL import Image, ImageOps, ImageEnhance, ImageFilter

# SQLAlchemy typing only (safe)
try:
    from sqlalchemy.orm import Session
except Exception:  # pragma: no cover
    Session = Any  # type: ignore
from io import BytesIO

try:
    from docx import Document  # python-docx
except Exception:
    Document = None


from .schema import (
    HealthResponse,
    OCRResponse,
    ClauseCheckRequest,
    ClauseCheckResponse,
    ClauseCheckResponseWithML,
)
from .utils_text import detect_language, norm_ar, normalize_for_rules, split_into_clauses
from .rag_rule_mapping import chunks_to_rule_ids

try:
    from .documentai_ocr import (
        documentai_extract_text,
        documentai_extract_pages,
        _HAS_DOCUMENTAI,
        _is_configured as _documentai_configured,
    )
except Exception:
    _HAS_DOCUMENTAI = False
    documentai_extract_text = lambda *a, **k: ""  # type: ignore
    documentai_extract_pages = lambda *a, **k: []  # type: ignore
    _documentai_configured = lambda: False  # type: ignore

# DocumentAI.py (bytes-based pipeline): preferred when configured
try:
    from .DocumentAI import (
        documentai_extract_pages_from_bytes,
        _is_configured as _documentai_py_configured,
    )
except Exception:
    documentai_extract_pages_from_bytes = None  # type: ignore
    _documentai_py_configured = lambda: False  # type: ignore

# ✅ cross-border detector
from .cross_border import detect_cross_border

# ============================================================
# ✅ Phase B: ML predictor import (SAFE)
# ============================================================
# ============================================================
# ✅ Phase B: ML predictor import (SAFE)
# ============================================================
# Why this exists:
# - Some versions of ml_predictor expose `predict_rule_scores_full`
# - Others expose `predict_rule_scores`
# We support both without breaking runtime.
try:
    from .ml_predictor import _HAS_ML_PREDICTOR  # type: ignore
    try:
        # preferred name (used by main.py below)
        from .ml_predictor import predict_rule_scores_full  # type: ignore
    except Exception:
        # fallback name
        from .ml_predictor import predict_rule_scores as predict_rule_scores_full  # type: ignore

    # Optional legacy helper (not strictly required, but kept if present)
    try:
        from .ml_predictor import predict_rule_ids  # type: ignore
    except Exception:
        predict_rule_ids = None  # type: ignore

except Exception as e:
    _HAS_ML_PREDICTOR = False
    predict_rule_ids = None  # type: ignore
    predict_rule_scores_full = None  # type: ignore
    print("[ML][WARN] ML predictor import failed:", repr(e))

# ============================================================
# ✅ Model-ML predictor (law-aware multilabel/binary; no rule engine)
# ============================================================
try:
    from .model_ml_predictor import (
        has_model_ml_predictor as _HAS_MODEL_ML_PREDICTOR,
        predict_rule_scores_full as model_ml_predict_rule_scores_full,
        predict_violation_risk_safe as model_ml_predict_violation_risk_safe,
        rule_scores_to_rule_hits as model_ml_rule_scores_to_rule_hits,
    )
except Exception as e:
    _HAS_MODEL_ML_PREDICTOR = False
    model_ml_predict_rule_scores_full = None  # type: ignore
    model_ml_predict_violation_risk_safe = None  # type: ignore
    model_ml_rule_scores_to_rule_hits = None  # type: ignore
    print("[ML][WARN] model_ml_predictor import failed:", repr(e))

# Legacy unified predictor (optional fallback)
try:
    from .unified_predictor import (
        has_unified_predictor as _HAS_UNIFIED_PREDICTOR,
        predict_violation_risk_safe as unified_predict_violation_risk_safe,
    )
except Exception:
    _HAS_UNIFIED_PREDICTOR = False
    unified_predict_violation_risk_safe = None  # type: ignore

# ============================================================
# ✅ Backend Enhancements: SAFE imports + DEBUG prints
# ============================================================
_HAS_DB = False
init_db = None

_HAS_AUTH = False
auth_router = None

_HAS_ANALYSES = False
analyses_router = None

_HAS_LOG_MW = False
RequestLoggingMiddleware = None

_HAS_LIMITER = False
limiter = None
SlowAPIMiddleware = None

# --- DB init (optional) ---
try:
    from app.db.init_db import init_db  # type: ignore
    _HAS_DB = True
except Exception as e:
    init_db = None
    _HAS_DB = False
    print("❌ Failed to import init_db:", repr(e))

# --- Auth router (optional) ---
try:
    from app.routers.auth import router as auth_router  # type: ignore
    _HAS_AUTH = True
except Exception as e:
    auth_router = None
    _HAS_AUTH = False
    print("❌ Failed to import auth_router:", repr(e))

# --- Analyses router (optional) ---
try:
    from app.routers.analyses import router as analyses_router  # type: ignore
    _HAS_ANALYSES = True
except Exception as e:
    analyses_router = None
    _HAS_ANALYSES = False
    print("❌ Failed to import analyses_router:", repr(e))

chat_router = None
try:
    from app.routers.chat import router as chat_router  # type: ignore
except Exception as e:
    chat_router = None
    print("❌ Failed to import chat_router:", repr(e))

# --- Request logging middleware (optional) ---
try:
    from app.middleware.request_logging import RequestLoggingMiddleware  # type: ignore
    _HAS_LOG_MW = True
except Exception as e:
    RequestLoggingMiddleware = None
    _HAS_LOG_MW = False
    print("❌ Failed to import RequestLoggingMiddleware:", repr(e))

# --- Rate limiting middleware (optional) ---
try:
    from app.middleware.rate_limit import limiter  # type: ignore
    from slowapi.middleware import SlowAPIMiddleware  # type: ignore
    _HAS_LIMITER = True
except Exception as e:
    limiter = None
    SlowAPIMiddleware = None
    _HAS_LIMITER = False
    print("❌ Failed to import rate limiter middleware:", repr(e))


# ============================================================
# ✅ DB/Auth deps (SAFE) — avoid crashes if modules missing
# ============================================================
_HAS_PERSIST = False
get_db = None
Analysis = None
User = Any
get_current_user = None
get_current_user_optional = None

try:
    from .db.session import get_db as _get_db  # type: ignore
    from .db.models import Analysis as _Analysis, User as _User  # type: ignore
    from .core.deps import get_current_user as _get_current_user  # type: ignore
    from fastapi.security import OAuth2PasswordBearer  # type: ignore

    get_db = _get_db
    Analysis = _Analysis
    User = _User
    get_current_user = _get_current_user
    _HAS_PERSIST = True

    # Optional auth dependency:
    # - Allows running /ocr_check_and_search without a JWT when save=false
    # - Still enforces auth when save=true
    _oauth2_optional = OAuth2PasswordBearer(tokenUrl="/auth/login/form", auto_error=False)

    def get_current_user_optional(  # type: ignore
        token: Optional[str] = Depends(_oauth2_optional),
        db: Session = Depends(_get_db),  # type: ignore
    ) -> Any:
        if not token:
            return None
        try:
            return _get_current_user(token=token, db=db)  # type: ignore
        except HTTPException:
            return None
except Exception as e:
    _HAS_PERSIST = False
    print("❌ Failed to import DB/Auth deps (get_db/models/deps):", repr(e))

    def _missing_db_dep() -> Generator[Any, None, None]:
        raise HTTPException(status_code=503, detail="DB layer not configured / missing imports.")

    def _missing_auth_dep() -> Any:
        raise HTTPException(status_code=503, detail="Auth layer not configured / missing imports.")

    get_db = _missing_db_dep  # type: ignore
    get_current_user = _missing_auth_dep  # type: ignore
    get_current_user_optional = lambda: None  # type: ignore


# ============================================================
# RAG imports (optional) + SAFE FALLBACKS
# ============================================================
try:
    from .rag_utils import Retriever  # type: ignore
except Exception:
    Retriever = None

# Prefer Legal Rag bridge (uses Legal Rag/src VectorStore, RAGEngine, DataIngestion)
try:
    from . import legal_rag_bridge as rag_chromadb  # type: ignore
except Exception:
    try:
        from . import rag_chromadb  # fallback: app ChromaDB over Legal RAG corpus
    except Exception:
        rag_chromadb = None  # type: ignore


def _fallback_load_chunks_as_docs(chunks_dir: Path) -> List[Dict[str, Any]]:
    docs: List[Dict[str, Any]] = []
    chunks_dir = Path(chunks_dir)
    if not chunks_dir.exists():
        return docs

    jsonl_files = sorted(chunks_dir.glob("*.jsonl"))
    for fp in jsonl_files:
        try:
            with fp.open("r", encoding="utf-8") as f:
                for _, line in enumerate(f, start=1):
                    line = line.strip()
                    if not line:
                        continue
                    try:
                        obj = json.loads(line)
                    except Exception:
                        continue

                    text = (obj.get("normalized_text") or obj.get("text") or obj.get("page_content") or "").strip()
                    if len(text) < 20:
                        continue

                    meta = {
                        "source_file": fp.name,
                        "id": obj.get("id"),
                        "law": obj.get("law"),
                        "article": obj.get("article"),
                        "title": obj.get("title"),
                        "source": obj.get("source") or fp.stem,
                        "page": obj.get("page"),
                        "chunk_id": obj.get("chunk_id"),
                    }
                    docs.append({"page_content": text, "metadata": meta})
        except Exception:
            continue

    return docs


class _FallbackRetriever:
    def __init__(self):
        self.docs: List[Dict[str, Any]] = []

    def build_index(self, docs: List[Dict[str, Any]]):
        self.docs = docs or []

    def _score(self, query: str, text: str) -> float:
        q = (query or "").strip()
        t = (text or "").strip()
        if not q or not t:
            return 0.0

        q_low = q.lower()
        t_low = t.lower()

        contains = 1.0 if q_low in t_low else 0.0

        import re
        q_tokens = set(re.findall(r"[a-zA-Z\u0600-\u06FF0-9]+", q_low))
        t_tokens = set(re.findall(r"[a-zA-Z\u0600-\u06FF0-9]+", t_low))
        overlap = 0.0
        if q_tokens and t_tokens:
            overlap = len(q_tokens & t_tokens) / max(1, len(q_tokens))

        import difflib
        fuzzy = difflib.SequenceMatcher(None, q_low, t_low[: min(len(t_low), 2000)]).ratio()

        return (2.0 * contains) + (1.2 * overlap) + (0.3 * fuzzy)

    def search(self, query: str, top_k: int = 5) -> List[Dict[str, Any]]:
        scored: List[Dict[str, Any]] = []
        for d in self.docs:
            text = d.get("page_content") or ""
            meta = d.get("metadata") or {}
            s = self._score(query, text)
            if s <= 0:
                continue
            scored.append({"score": float(s), "text": text, "metadata": meta})

        scored.sort(key=lambda x: x["score"], reverse=True)
        return scored[: max(1, int(top_k))]


try:
    from .chunks_loader import load_chunks_as_docs  # type: ignore
except Exception:
    load_chunks_as_docs = _fallback_load_chunks_as_docs

if Retriever is None:
    Retriever = _FallbackRetriever  # type: ignore


# ============================================================
# FastAPI instance
# ============================================================
api = FastAPI(title="legalai")
app = api  # alias for uvicorn

# ============================================================
# CORS — allow frontend (e.g. localhost:5173) to call the API
# ============================================================
api.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://127.0.0.1:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ============================================================
# ✅ Middleware + Routers (SAFE)
# ============================================================
if _HAS_LOG_MW and RequestLoggingMiddleware:
    api.add_middleware(RequestLoggingMiddleware)

if _HAS_LIMITER and limiter and SlowAPIMiddleware:
    api.state.limiter = limiter
    api.add_middleware(SlowAPIMiddleware)

if _HAS_AUTH and auth_router:
    api.include_router(auth_router)
    print("✅ Auth router mounted:", getattr(auth_router, "prefix", None))

if _HAS_ANALYSES and analyses_router:
    api.include_router(analyses_router)
    print("✅ Analyses router mounted:", getattr(analyses_router, "prefix", None))

if chat_router:
    api.include_router(chat_router)
    print("✅ Chat router mounted:", getattr(chat_router, "prefix", None))

# ============================================================
# Optional jobs router (SAFE – stub included; mount only if present)
# ============================================================
try:
    from app.jobs import router as jobs_router
    api.include_router(jobs_router)
    print("✅ Jobs router mounted:", getattr(jobs_router, "prefix", None))
except Exception as e:
    print("ℹ️ Jobs router not mounted:", repr(e))


# ============================================================
# Paths + Rule engine (source of truth when available)
# ============================================================
BASE_DIR = Path(__file__).resolve().parent.parent
RULES_DIR = BASE_DIR / "rules"
LAWS_DIR = BASE_DIR / "laws"
CHUNKS_DIR = BASE_DIR / "chunks"

retriever: Optional[Any] = None
rule_engine: Optional[Any] = None

try:
    from .rules import RuleEngine
    if RULES_DIR.exists() and LAWS_DIR.exists():
        rule_engine = RuleEngine(RULES_DIR, LAWS_DIR)
except Exception:
    rule_engine = None


# ============================================================
# 0) Root + Health
# ============================================================
@api.get("/")
def root():
    return {"name": "legalai", "status": "ok"}


@api.get("/health", response_model=HealthResponse)
def health() -> HealthResponse:
    return HealthResponse()


# ============================================================
# 1) OCR helpers (IMPROVED)
# ============================================================
def _preprocess_for_ocr(img: Image.Image) -> Image.Image:
    try:
        img = img.convert("L")
        img = ImageOps.autocontrast(img)
        img = img.filter(ImageFilter.SHARPEN)
        img = ImageEnhance.Contrast(img).enhance(1.4)
        img = img.point(lambda x: 0 if x < 160 else 255, mode="1")
        img = img.convert("L")
        return img
    except Exception:
        return img


def _docx_bytes_to_text(data: bytes) -> str:
    try:
        from docx import Document
        import io

        doc = Document(io.BytesIO(data))

        parts: List[str] = []

        # 1) paragraphs
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)

        # 2) tables (VERY IMPORTANT for contracts)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = (cell.text or "").strip()
                    if t:
                        parts.append(t)

        text = "\n".join(parts).strip()
        return text
    except Exception as e:
        print("❌ DOCX extract failed:", repr(e))
        return ""


def _docx_images_to_text(data: bytes) -> str:
    """
    If a DOCX is scanned (images), extract images from word/media and OCR them.
    Deterministic fallback, no schema change.
    """
    import io, zipfile

    texts: List[str] = []
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            media_files = [n for n in z.namelist() if n.startswith("word/media/")]
            for name in sorted(media_files):
                try:
                    img_bytes = z.read(name)
                    t = _ocr_image_to_text(img_bytes) or ""
                    t = (t or "").strip()
                    if t:
                        texts.append(t)
                except Exception:
                    continue
    except Exception as e:
        print("❌ DOCX image OCR failed:", repr(e))
        return ""

    return "\n".join(texts).strip()


def _ocr_image_to_text(img_bytes: bytes) -> str:
    try:
        import pytesseract
        img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
        img = _preprocess_for_ocr(img)
        config = "--oem 1 --psm 6"
        return pytesseract.image_to_string(img, lang="ara+eng", config=config)
    except Exception:
        return ""


def _is_text_too_short(txt: str) -> bool:
    return len((txt or "").strip()) < 40


def _normalize_contract_text(raw_text: str) -> str:
    if not raw_text:
        return ""
    t = norm_ar(raw_text)
    t = t.replace("ًل", "لا").replace("اًل", "ال")
    fixes = [
        ("جميه", "جنيه"),
        ("جنيه مصي", "جنيه مصري"),
        ("مصي", "مصري"),
        ("ديب", "دبي"),
        ("الامارا", "الامارات"),
        ("الامارا ت", "الامارات"),
        ("الامــارا", "الامارات"),
        ("الامارات العربي املتحدة", "الامارات العربية المتحدة"),
        ("االمارات", "الامارات"),
        ("العربي املتحدة", "العربية المتحدة"),
        ("قانو ن", "قانون"),
        ("القانو ن", "القانون"),
        ("ة مد", "مدة"),
        ("ا هذ لعقد ا", "هذا العقد"),
        ("ا هذ", "هذا"),
        ("لعا قد", "لعقد"),
        ("فوة", "فترة"),
        ("اختبﺎر", "اختبار"),
        ("ا نذار", "انذار"),
        ("بﺄجر", "باجر"),
    ]
    for a, b in fixes:
        t = t.replace(a, b)

    return t.strip()


# ============================================================
# OCR endpoint (unchanged schema)
# ============================================================
@api.post("/ocr", response_model=OCRResponse)
async def ocr(file: UploadFile = File(...)) -> OCRResponse:
    data = await file.read()
    text = ""

    is_pdf = (
        file.filename.lower().endswith(".pdf")
        or (file.content_type or "").lower() == "application/pdf"
    )

    if is_pdf:
        if _documentai_configured() and documentai_extract_text:
            try:
                dai_text = documentai_extract_text(data)
                if dai_text and len(dai_text.strip()) > 30:
                    text = dai_text
                else:
                    raise ValueError("Document AI returned insufficient text")
            except Exception:
                pass
        if not text:
            try:
                doc = fitz.open(stream=data, filetype="pdf")
                parts: List[str] = []
                for page in doc:
                    txt = page.get_text("text") or ""
                    if _is_text_too_short(txt):
                        pix = page.get_pixmap(dpi=300)
                        txt = _ocr_image_to_text(pix.tobytes("png"))
                    parts.append(txt)
                text = "\n".join(parts)
            except Exception:
                text = _ocr_image_to_text(data)
    else:
        text = _ocr_image_to_text(data)

    return OCRResponse(text=text)


# ============================================================
# 2) Match/hits output normalizer (legacy; model_ML returns structured hits directly)
# ============================================================
def _coerce_hits(raw: Any) -> List[Dict[str, Any]]:
    if raw is None:
        return []
    if isinstance(raw, list):
        return [x if isinstance(x, dict) else {"description": str(x)} for x in raw]
    if isinstance(raw, dict):
        if "matches" in raw and isinstance(raw["matches"], list):
            return [h if isinstance(h, dict) else {"description": str(h)} for h in raw["matches"]]
        if "hits" in raw:
            h = raw["hits"]
            if isinstance(h, list):
                return [x if isinstance(x, dict) else {"description": str(x)} for x in h]
            if isinstance(h, dict):
                return [v if isinstance(v, dict) else {"description": str(v)} for v in h.values()]
        return [raw]
    return [{"description": str(raw)}]


# ============================================================
# 2.5) Summary post-processing tweaks
# ============================================================
def _apply_leave_waiver_to_summary(
    labor_summary: Dict[str, Any],
    rule_hits: List[Dict[str, Any]]
) -> Dict[str, Any]:

    if not labor_summary or not isinstance(labor_summary, dict):
        return labor_summary

    annual = labor_summary.get("annual_leave")
    if not isinstance(annual, dict):
        return labor_summary

    waiver_hit = next(
        (h for h in (rule_hits or []) if h.get("rule_id") == "LABOR25_ANNUAL_LEAVE_WAIVER"),
        None
    )
    deferral_hit = next(
        (h for h in (rule_hits or []) if h.get("rule_id") == "LABOR25_ANNUAL_LEAVE_DEFERRAL"),
        None
    )

    if not waiver_hit and not deferral_hit:
        return labor_summary

    # ✅ safe: choose text from whichever hit exists
    base_hit = waiver_hit or deferral_hit
    hit_text = (base_hit.get("matched_text") or "").strip()
    hit_text_low = hit_text.lower()

    deferral_markers = [
        "السنة الأولى", "السنه الاولى", "اول سنة", "أول سنة", "خلال اول", "خلال أول",
        "بعد سنة", "بعد مرور سنة", "بعد مرور عام", "بعد عام",
        "first year", "after one year", "after a year",
    ]

    looks_like_deferral = any(m.lower() in hit_text_low for m in deferral_markers)

    # CASE A: deferral (explicit rule OR text indicates first-year/deferral)
    if deferral_hit or looks_like_deferral:
        annual["present"] = True
        annual["min_violation"] = False
        annual["status"] = "ok"
        annual["message"] = (
            "تم تنظيم/تأجيل استحقاق الإجازة السنوية (مثل السنة الأولى) وليس إسقاطًا دائمًا للحق. "
            "| Annual leave is deferred/organized (e.g., first-year policy), not a permanent waiver."
        )
        labor_summary["annual_leave"] = annual
        return labor_summary

    # CASE B: real waiver
    if waiver_hit:
        annual["present"] = True
        annual["min_violation"] = True
        annual["status"] = "violation"
        annual["message"] = (
            "يوجد شرط يسقط/يلغي حق الإجازة السنوية أو يمنع العامل منها بشكل دائم (مخالفة). "
            "| Permanent annual leave waiver/forfeiture detected (violation)."
        )
        labor_summary["annual_leave"] = annual
        return labor_summary

    return labor_summary


# ============================================================
# ✅ RAG helpers: query boosting + safe search wrapper + filtering
# ============================================================
def build_rag_query(
    user_query: str,
    contract_text: str,
    tags: List[str],
    rule_texts: Optional[List[str]] = None,
) -> str:
    """
    Boost query toward relevant Labor Law contract articles, and reduce generic hits.
    If rule_texts is provided, append rule descriptions/rationales to favor law chunks relevant to those rules.
    """
    q = norm_ar(user_query or "")
    _ = norm_ar(contract_text or "")

    anchor = "قانون العمل رقم 14 لسنة 2025 عقد العمل الفردي يجب أن يتضمن العقد"

    extra: List[str] = []
    q_low = q.lower()

    if any(k in q_low for k in ["اجر", "الأجر", "مرتب", "راتب", "بدلات", "حوافز", "مزايا", "صافي", "إجمالي"]):
        extra.append("الأجر طريقة الصرف موعد صرف الأجر البدلات المزايا المالية الاستقطاعات التأمينات الضرائب")

    if any(k in q_low for k in ["صاحب العمل", "الشركة", "طرف اول", "الطرف الأول", "العنوان", "مقر", "محل العمل"]):
        extra.append("اسم صاحب العمل عنوان محل العمل مقر ممارسة النشاط بيانات صاحب العمل")

    if any(k in q_low for k in ["العامل", "الموظف", "طرف تاني", "الطرف الثاني", "رقم قومي", "التأمين", "محل الإقامة", "مهنة", "وظيفة"]):
        extra.append("اسم العامل مهنته الرقم القومي الرقم التأميني محل الإقامة بيانات العامل")

    if any(k in q_low for k in ["مدة", "محدد", "غير محدد", "فترة اختبار", "تجديد", "بداية", "نهاية"]):
        extra.append("مدة العقد عقد محدد المدة عقد غير محدد المدة تاريخ بداية العقد تاريخ نهاية العقد فترة الاختبار تجديد العقد")

    if any(k in q_low for k in ["انهاء", "فصل", "استقالة", "اخطار", "إنذار", "تعويض", "إنهاء"]):
        extra.append("إنهاء عقد العمل مهلة الإخطار الفصل التعسفي التعويض")

    if "cross_border" in (tags or []):
        extra.append("تشغيل المصريين بالخارج التعاقدات مع جهة أجنبية")

    if rule_texts:
        joined = " ".join(t for t in rule_texts if t and isinstance(t, str))[:2000]
        if joined:
            extra.append(joined)

    boosted = " ".join([q, anchor] + extra).strip()
    return norm_ar(boosted)


def _rag_search_safe(
    query_text: str,
    top_k: int = 5,
    law_filters: Optional[Dict[str, Any]] = None,
    min_score: float = 0.0,
) -> List[Dict[str, Any]]:
    """
    Prefer ChromaDB (Legal RAG) when available; else retriever (FAISS/hash).
    Same return shape: [{"score", "text", "metadata"}, ...]
    """
    if not query_text or not query_text.strip():
        return []

    # Prefer ChromaDB (Legal RAG corpus)
    if rag_chromadb is not None and getattr(rag_chromadb, "is_available", lambda: False)():
        try:
            hits = rag_chromadb.search(
                query=query_text,
                top_k=top_k,
                filters=law_filters,
                min_score=min_score,
            )
            if hits is not None:
                return hits
        except Exception:
            pass

    if not retriever:
        return []

    try:
        return retriever.search(query_text, top_k=top_k, filters=law_filters, min_score=min_score)  # type: ignore
    except TypeError:
        try:
            return retriever.search(query_text, top_k=top_k)  # type: ignore
        except Exception:
            return []
    except Exception:
        return []


def _filter_to_labor_only(hits: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    for h in hits or []:
        md = h.get("metadata") or {}
        law = str(md.get("law") or "")
        src = str(md.get("source") or "")
        if ("قانون العمل" in law) or (src == "labor14_2025"):
            out.append(h)
    return out


def _dedupe_rag_blocks_by_rule(rag_by_violation: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    seen = set()
    out: List[Dict[str, Any]] = []
    for b in rag_by_violation or []:
        rid = b.get("rule_id")
        if rid in seen:
            continue
        seen.add(rid)
        out.append(b)
    return out


def _rule_evidence(rule_hits: List[Dict[str, Any]], rule_ids: List[str]) -> Dict[str, Any]:
    evidence = {}
    for rid in rule_ids:
        hit = next((h for h in rule_hits if h.get("rule_id") == rid), None)
        evidence[rid] = {
            "hit": hit is not None,
            "matched_text": hit.get("matched_text") if hit else None,
            "chunk_id": hit.get("chunk_id") if hit else None,
            "severity": hit.get("severity") if hit else None,
        }
    return evidence


# ============================================================
# 3) Clause-level check  ✅ REPLACED بالكامل (ML shortlist + fallback)
# ============================================================
@api.post("/check_clause", response_model=ClauseCheckResponseWithML)
def check_clause(req: ClauseCheckRequest):
    try:
        text = req.clause_text or ""
        lang = (getattr(req, "language", None) or detect_language(text)).lower()
        text_norm = normalize_for_rules(text) if lang == "ar" else text.strip()

        law_scope = req.law_scope or ["labor"]

        # ----------------------------
        # ✅ Model-ML only (no rule engine): violations from model_ML bundles
        # ----------------------------
        ml_used = False
        ml_predictions = None
        deduped: List[Dict[str, Any]] = []
        unified_risk = None
        unified_above = None

        if _HAS_MODEL_ML_PREDICTOR and model_ml_predict_rule_scores_full is not None and model_ml_rule_scores_to_rule_hits is not None:
            try:
                preds = model_ml_predict_rule_scores_full(text_norm, sort=True, use_law_retrieval=True)
                ml_predictions = preds
                ml_used = True
                deduped = model_ml_rule_scores_to_rule_hits(preds)
                # Ensure each hit has rule_id / id for response compatibility
                for h in deduped:
                    if h.get("rule_id") and not h.get("id"):
                        h["id"] = h["rule_id"]
            except Exception as e:
                print("❌ Model-ML predict error:", repr(e))
                ml_used = False
                ml_predictions = None
                deduped = []

        if _HAS_MODEL_ML_PREDICTOR and model_ml_predict_violation_risk_safe is not None:
            try:
                res = model_ml_predict_violation_risk_safe(text_norm, use_law_at_inference=True)
                if res is not None:
                    unified_risk, unified_above = res
            except Exception:
                pass
        elif _HAS_UNIFIED_PREDICTOR and unified_predict_violation_risk_safe is not None:
            try:
                res = unified_predict_violation_risk_safe(text_norm, use_law_at_inference=False)
                if res is not None:
                    unified_risk, unified_above = res
            except Exception:
                pass

        resp = ClauseCheckResponseWithML(
            clause_text=text,
            language=lang,
            matches=deduped,
            ml_used=ml_used,
            ml_predictions=ml_predictions,
            unified_ml_risk=unified_risk,
            unified_ml_above_threshold=unified_above,
        )
        return JSONResponse(content=resp.model_dump(), media_type="application/json; charset=utf-8")

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Rule check failed: {e.__class__.__name__}: {e}")

# Single consolidated startup report (printed once per process, e.g. avoid duplicate on reload)
_startup_report_done = False


# ============================================================
# 4) Startup – DB init + RAG init + health checks
# ============================================================
@api.on_event("startup")
def startup():
    global retriever, _startup_report_done

    # DB init
    db_ok = False
    if _HAS_DB and init_db:
        try:
            init_db()
            db_ok = True
            print("✅ DB initialized (SQLite).")
        except Exception as e:
            print("❌ Failed to init DB:", repr(e))
    else:
        print("ℹ️ DB not configured (optional).")

    # RAG init: ChromaDB (Legal RAG) preferred; in-memory retriever fallback
    chroma_count = 0
    chroma_name = None
    device_used = None
    embedding_model = None
    artifacts_ok = False
    if rag_chromadb is not None and getattr(rag_chromadb, "is_available", lambda: False)():
        try:
            chroma_count = getattr(rag_chromadb, "get_collection_count", lambda: 0)()
            chroma_name = getattr(rag_chromadb, "get_collection_name", lambda: None)()
            device_used = getattr(rag_chromadb, "get_device", lambda: None)()
            embedding_model = getattr(rag_chromadb, "get_embedding_model_name", lambda: None)()
            artifacts_ok = True
            print(f"✅ ChromaDB RAG ready (Legal RAG corpus): {chroma_count} docs.")
        except Exception as e:
            print("❌ ChromaDB RAG init failed:", repr(e))
    else:
        print("ℹ️ ChromaDB RAG not available (optional).")

    retriever_docs = 0
    try:
        docs = load_chunks_as_docs(CHUNKS_DIR)
        retriever = Retriever()  # type: ignore
        retriever.build_index(docs)
        retriever_docs = len(docs)
        print(f"✅ Retriever initialized with {retriever_docs} docs.")
    except Exception as e:
        retriever = None
        print("❌ Failed to init Retriever:", repr(e))

    # Single consolidated startup report (once per process)
    if not _startup_report_done:
        print(
            "[Startup] RAG: Chroma (Legal RAG)=%s docs (collection=%s); in-memory Retriever (chunks)=%s docs. "
            "DB=%s, artifacts=%s, device=%s, embedding_model=%s"
            % (
                chroma_count,
                chroma_name or "n/a",
                retriever_docs,
                "ok" if db_ok else "n/a",
                "ok" if artifacts_ok else "n/a",
                device_used or "n/a",
                embedding_model or "n/a",
            )
        )
        _startup_report_done = True


@api.on_event("shutdown")
def on_shutdown():
    """Graceful shutdown: cleanup and avoid noisy tracebacks on reload/Ctrl+C."""
    try:
        print("Shutting down gracefully.")
    except Exception:
        pass


# ============================================================
# 5) Labor summary helper (ML-based: rule_hits are model_ML violations)
# ============================================================
def _build_labor_summary(rule_hits: List[Dict[str, Any]]) -> Dict[str, Any]:
    def bi(ar: str, en: str) -> str:
        return f"{ar} | {en}"

    ml_violations = [
        {"rule_id": h.get("rule_id"), "score": h.get("score"), "description": h.get("description") or ""}
        for h in (rule_hits or [])
        if h.get("rule_id")
    ]
    has_violations = len(ml_violations) > 0
    status = "violations_detected" if has_violations else "ok"
    message = bi(
        "تم اكتشاف مخالفات محتملة (نموذج ML). راجع ml_violations." if has_violations else "لم يتم اكتشاف مخالفات من نموذج ML.",
        "Potential violations detected (ML model). See ml_violations." if has_violations else "No violations detected by ML model.",
    )

    summary: Dict[str, Any] = {
        "status": status,
        "message": message,
        "ml_violations": ml_violations,
        "source": "model_ml",
    }
    # Minimal backward-compatible keys (consumers may expect these)
    summary["employer_info"] = {"present": None, "status": "ml_based", "message": bi("التقييم من نموذج ML.", "Assessment from ML model.")}
    summary["employee_info"] = {"present": None, "status": "ml_based", "message": bi("التقييم من نموذج ML.", "Assessment from ML model.")}
    summary["salary"] = {"present": None, "status": "ml_based", "message": bi("التقييم من نموذج ML.", "Assessment from ML model.")}
    summary["contract_duration"] = {"present": None, "status": "ml_based", "message": bi("التقييم من نموذج ML.", "Assessment from ML model.")}
    summary["working_hours"] = {"present": None, "status": "ml_based", "message": bi("التقييم من نموذج ML.", "Assessment from ML model.")}
    summary["annual_leave"] = {"present": None, "min_violation": None, "status": "ml_based", "message": bi("التقييم من نموذج ML.", "Assessment from ML model.")}
    summary["probation"] = {"present": None, "limit_violation": None, "status": "ml_based", "message": bi("التقييم من نموذج ML.", "Assessment from ML model.")}
    return summary


# ============================================================
# 5.5) Cross-border summary helper (ML-based: rule_hits are model_ML violations)
# ============================================================
def _build_cross_border_summary(
    rule_hits: List[Dict[str, Any]],
    contract_tags: List[str],
    cb_info: Dict[str, Any],
) -> Dict[str, Any]:
    def bi(ar: str, en: str) -> str:
        return f"{ar} | {en}"

    is_cross_border = ("cross_border" in (contract_tags or [])) or bool(cb_info.get("enabled"))

    if not is_cross_border:
        return {
            "enabled": False,
            "status": "not_applicable",
            "message": bi(
                "هذا العقد لا يبدو عقد عمل دولي/خارج مصر.",
                "This contract does not appear to be cross-border / outside Egypt.",
            ),
        }

    ml_violations = [
        {"rule_id": h.get("rule_id"), "score": h.get("score"), "description": h.get("description") or ""}
        for h in (rule_hits or [])
        if h.get("rule_id")
    ]
    summary: Dict[str, Any] = {
        "enabled": True,
        "status": "ok" if not ml_violations else "violations_detected",
        "message": bi("ملخص العقد الدولي (تقييم من نموذج ML).", "Cross-border contract summary (ML-based)."),
        "signals": cb_info.get("matches", []),
        "ml_violations": ml_violations,
        "source": "model_ml",
    }
    return summary


def _labor_applicability(full_text: str, cb: Dict[str, Any], contract_tags: List[str]) -> Dict[str, Any]:
    """
    Decide whether Egyptian Labor Law scope should be applied.

    Policy:
    - Cross-border alone does NOT disable Egyptian labor law.
    - Disable Egyptian labor law ONLY if there are strong FOREIGN governing-law/jurisdiction signals
      AND there is NO strong Egyptian governing-law signal.
    - If Egyptian governing law is explicitly stated => keep labor applicable even if cross-border.
    """
    t = (full_text or "")
    t_low = t.lower()

    is_cb = bool(cb.get("enabled")) or ("cross_border" in (contract_tags or []))

    # Strong Egypt governing law / jurisdiction signals (keep labor applicable)
    egypt_strong = [
        "القانون المصري",
        "قانون العمل المصري",
        "يخضع لأحكام القانون المصري",
        "يخضع للقانون المصري",
        "تطبق أحكام القانون المصري",
        "محاكم مصر",
        "المحاكم المصرية",
        "محكمة مصر",
        "جمهورية مصر العربية",
        "القاهرة",
        "الجيزة",
        "الاسكندرية",
        "egyptian law",
        "laws of egypt",
        "courts of egypt",
        "egypt",
    ]
    egypt_hit = any(s.lower() in t_low for s in egypt_strong)

    # Strong UAE/foreign governing law/jurisdiction signals (disable labor when no egypt_hit)
    # IMPORTANT: must be truly "strong" (country-specific), not generic words like "طبقا للقانون" or "محاكم"
    foreign_strong = [
        # =========================
        # UAE specific (strong)
        # =========================
        "قانون العمل الإماراتي",
        "المرسوم بقانون اتحادي",
        "مرسوم بقانون اتحادي",
        "اتحادي رقم 33 لسنة 2021",
        "33 لسنة 2021",
        "mohre",
        "وزارة الموارد البشرية والتوطين",
        "التوطين",
        "محاكم الإمارات",
        "المحاكم الإماراتية",
        "دولة الإمارات",
        "الإمارات العربية المتحدة",
        "united arab emirates",
        "uae",
        "dubai",
        "دبي",
        "أبوظبي",
        "ابوظبي",
        "داخل دولة الإمارات",
        "داخل أراضي دولة الإمارات",

        # =========================
        # Strong English governing law clauses (country-specific)
        # =========================
        "this agreement shall be governed by the laws of uae",
        "this agreement shall be governed by the laws of the united arab emirates",
        "governed by the laws of the united arab emirates",
        "subject to the laws of the united arab emirates",
        "courts of dubai",
        "courts of the united arab emirates",

        # =========================
        # Strong Arabic governing law clauses with a country explicitly
        # =========================
        "يخضع هذا العقد لقانون دولة الإمارات",
        "يخضع هذا العقد لقانون الإمارات",
        "يخضع لقانون دولة الإمارات",
        "يخضع لقانون الإمارات",
        "تختص محاكم الإمارات",
        "تختص المحاكم الإماراتية",
    ]

    foreign_hit = any(s.lower() in t_low for s in foreign_strong)

    # If explicitly Egyptian law => applicable even if cross-border
    if egypt_hit:
        return {
            "applicable": True,
            "status": "applicable",
            "reason": "Egyptian governing law/jurisdiction explicitly detected",
            "jurisdiction": "EG",
            "is_cross_border": is_cb,
        }

    # If cross-border + strong foreign law/jurisdiction => not applicable
    if is_cb and foreign_hit:
        return {
            "applicable": False,
            "status": "not_applicable",
            "reason": "Strong foreign governing law/jurisdiction detected (no explicit Egyptian law)",
            "jurisdiction": "FOREIGN",
            "is_cross_border": is_cb,
        }

    # Default: keep applicable (even if cross-border) unless strong foreign says otherwise
    return {
        "applicable": True,
        "status": "applicable",
        "reason": "No strong foreign governing law override detected",
        "jurisdiction": "EG",
        "is_cross_border": is_cb,
    }


# ============================================================
# 6) OCR + Rules + RAG endpoint (IMPROVED)
# ============================================================
# Prefer llm folder (llm/generate.py) for explanation; fallback to app/local_llm
try:
    from llm.generate import explain_violation as _llm_explain_violation
    from llm.generate import is_available as _llm_is_available
except Exception:
    _llm_explain_violation = None  # type: ignore
    _llm_is_available = lambda: False  # type: ignore
try:
    from . import local_llm as _local_llm_module
except Exception:
    _local_llm_module = None  # type: ignore


@api.post("/ocr_check_and_search")
async def ocr_check_and_search(
    file: UploadFile = File(...),
    query: Optional[str] = Form(None),
    use_rag: Optional[bool] = Form(True),
    use_rag_for_rules: Optional[bool] = Form(True),  # RAG-first: run only rules matching RAG chunks
    use_ml: Optional[bool] = Form(True),  # ML fallback when RAG returns nothing
    use_llm: bool = Form(False),  # Local LLM explanation for violations (LFM2.5-1.2B-Instruct)
    llm_top_k: int = Form(2),  # limit number of violations explained by LLM
    llm_max_new_tokens: int = Form(200),  # cap tokens for faster inference
    save: bool = Form(False),
    db: Session = Depends(get_db),  # type: ignore
    current_user: Any = Depends(get_current_user_optional),  # type: ignore
):
    try:
        import numpy as np
        np.random.seed(42)  # reproducibility: same contract -> same score
        data = await file.read()

        # ---------- detect file types ----------
        filename = (file.filename or "").lower()
        ctype = (file.content_type or "").lower()

        is_docx = filename.endswith(".docx") or ("wordprocessingml" in ctype)
        is_pdf = filename.endswith(".pdf") or (ctype == "application/pdf")

        ocr_chunks: List[Dict[str, Any]] = []
        ocr_used_flag = False
        txt = ""  # prevents UnboundLocalError

        # ============================================================
        # 1) Text extraction / OCR -> build ocr_chunks
        # ============================================================
        if is_docx:
            txt = _docx_bytes_to_text(data) or ""
            ocr_used_flag = False

            # base chunk
            ocr_chunks = [{
                "id": "page_0",
                "page": 0,
                "text": txt,
                "left": None,
                "top": None,
                "width": None,
                "height": None,
                "conf": None,
            }]

            # If empty -> try OCR from embedded images (scanned DOCX)
            if not (txt or "").strip():
                txt2 = _docx_images_to_text(data) or ""
                if txt2.strip():
                    ocr_chunks[0]["text"] = txt2
                    ocr_used_flag = True
                    txt = txt2
                else:
                    raise HTTPException(
                        status_code=400,
                        detail="DOCX text extraction returned empty. The document may be scanned or content is not readable."
                    )

        elif is_pdf:
            # Prefer DocumentAI.py (bytes-based) when configured; else documentai_ocr
            dai_pages = []
            if documentai_extract_pages_from_bytes and _documentai_py_configured():
                try:
                    dai_pages = documentai_extract_pages_from_bytes(data)
                except Exception:
                    pass
            if not dai_pages and _documentai_configured() and documentai_extract_pages:
                try:
                    dai_pages = documentai_extract_pages(data)
                except Exception:
                    pass
            if dai_pages and any((p.get("text") or "").strip() for p in dai_pages):
                ocr_used_flag = True
                for p in dai_pages:
                    ocr_chunks.append({
                        "id": f"page_{p.get('page', len(ocr_chunks))}",
                        "page": p.get("page", len(ocr_chunks)),
                        "text": p.get("text", "") or "",
                        "left": None,
                        "top": None,
                        "width": None,
                        "height": None,
                        "conf": None,
                    })
            if not ocr_chunks:
                doc = fitz.open(stream=data, filetype="pdf")
                try:
                    for i, page in enumerate(doc):
                        page_txt = page.get_text("text") or ""
                        if _is_text_too_short(page_txt):
                            ocr_used_flag = True
                            pix = page.get_pixmap(dpi=300)
                            page_txt = _ocr_image_to_text(pix.tobytes("png")) or ""

                        ocr_chunks.append({
                            "id": f"page_{i}",
                            "page": i,
                            "text": page_txt,
                            "left": None,
                            "top": None,
                            "width": None,
                            "height": None,
                            "conf": None,
                        })
                finally:
                    doc.close()

        else:
            # image or unknown -> OCR
            ocr_used_flag = True
            txt = _ocr_image_to_text(data) or ""
            ocr_chunks = [{
                "id": "page_0",
                "page": 0,
                "text": txt,
                "left": None,
                "top": None,
                "width": None,
                "height": None,
                "conf": None,
            }]

        # ============================================================
        # 2) Normalize + confidence (OUTSIDE branches)
        # ============================================================
        for c in ocr_chunks:
            raw_text = c.get("text", "") or ""
            c["normalized_text"] = _normalize_contract_text(raw_text) if raw_text else ""

        for c in ocr_chunks:
            c["confidence"] = c["conf"] if isinstance(c.get("conf"), (int, float)) else None

        full_text = "\n\n".join((c.get("normalized_text") or "") for c in ocr_chunks)

        full_text_rules = normalize_for_rules(full_text)

        # Clause splitting (for linking rule hits to contract clauses)
        clause_spans = split_into_clauses(full_text, min_clause_len=15)
        clauses_for_response = [
            {"start": s, "end": e, "text": t[:500]} for s, e, t in clause_spans
        ]
        # Model-ML / unified: clause-level violation risk per clause
        clause_level_unified_risks: List[Dict[str, Any]] = []
        full_text_unified_risk: Optional[float] = None
        _predict_risk = model_ml_predict_violation_risk_safe if (_HAS_MODEL_ML_PREDICTOR and model_ml_predict_violation_risk_safe) else (unified_predict_violation_risk_safe if (_HAS_UNIFIED_PREDICTOR and unified_predict_violation_risk_safe) else None)
        _use_law = True  # model_ML uses law retrieval; unified typically False
        if _HAS_MODEL_ML_PREDICTOR and model_ml_predict_violation_risk_safe:
            _use_law = True
        elif _HAS_UNIFIED_PREDICTOR and unified_predict_violation_risk_safe:
            _use_law = False
        if _predict_risk is not None:
            try:
                for s, e, t in clause_spans:
                    if len(t.strip()) < 15:
                        continue
                    t_norm = normalize_for_rules(t)
                    res = _predict_risk(t_norm, use_law_at_inference=_use_law)
                    if res is not None:
                        score, above = res
                        clause_level_unified_risks.append({
                            "start": s, "end": e,
                            "text_preview": t[:200],
                            "unified_ml_risk": round(score, 4),
                            "unified_ml_above_threshold": above,
                        })
                res_full = _predict_risk(full_text_rules[:15000], use_law_at_inference=_use_law)
                if res_full is not None:
                    full_text_unified_risk = round(res_full[0], 4)
            except Exception:
                pass
        # ============================================================
        # 2.5) Contract type & tags
        # ============================================================
        def detect_contract_type(text: str) -> str:
            t = (text or "").lower()
            if any(x in t for x in ["remote", "work from home", "عن بعد", "عن بُعد", "عمل عن بعد"]):
                return "remote"
            if any(x in t for x in ["part time", "part-time", "دوام جزئي", "عمل جزئي", "ساعات عمل مرنة"]):
                return "part_time"
            return "full_time"

        def detect_contract_tags(text: str) -> List[str]:
            t = (text or "").lower()
            tags: List[str] = []
            import re

            strong_any = [
                r"\biban\b", r"\bswift\b", r"\bwise\b", r"\bpayoneer\b",
                r"\busd\b", r"\beur\b", r"\bgbp\b", r"\baed\b", r"\bsar\b",
                r"دولار", r"يورو", r"درهم", r"ريال", r"استرليني",
                r"bank\s*transfer", r"wire\s*transfer", r"تحويل\s*بنكي",
                r"mohre", r"وزارة\s*الموارد\s*البشرية", r"التوطين",
            ]
            geo_any = [
                r"\buae\b", r"united\s*arab\s*emirates",
                r"الامارات", r"الاماره", r"الامارا", r"الإمارات", r"الامارا.?ت",
                r"دبي", r"ديب",
                r"خارج\s*مصر", r"دولية", r"بالخارج", r"international", r"abroad", r"outside\s*egypt",
                r"saudi", r"\bksa\b", r"qatar", r"kuwait", r"germany", r"\buk\b", r"\busa\b", r"canada", r"europe",
                r"السعودية", r"قطر", r"الكويت", r"ألمانيا", r"اوروبا", r"أوروبا", r"انجلترا", r"أمريكا",
            ]
            weak_any = [
                r"fees", r"transfer\s*fees", r"رسوم\s*التحويل",
                r"tax", r"taxes", r"ضرائب", r"تأمينات", r"social\s*insurance", r"تأمين\s*اجتماعي",
            ]

            score = 0
            if any(re.search(p, t) for p in strong_any):
                score += 2
            if any(re.search(p, t) for p in geo_any):
                score += 1
            if any(re.search(p, t) for p in weak_any):
                score += 1

            if score >= 2:
                tags.append("cross_border")
            return tags

        contract_type = detect_contract_type(full_text)
        contract_tags = detect_contract_tags(full_text)

        cb = detect_cross_border(full_text)

        # Heuristic fallback
        if not cb.get("enabled"):
            import re
            if re.search(r"(الامارات|الاماره|الامارا|الإمارات|uae|united\s*arab\s*emirates)", full_text.lower()) or re.search(
                r"(دبي|ديب|dubai)", full_text.lower()
            ):
                cb = {"enabled": True, "reason": "heuristic_geo", "matches": ["geo_signal_detected"]}

        is_cb = bool(cb.get("enabled"))
        if is_cb and "cross_border" not in contract_tags:
            contract_tags.append("cross_border")

        # Labor applicability gate
        labor_appx = _labor_applicability(full_text, cb, contract_tags)
        labor_applicable = bool(labor_appx.get("applicable"))

        rag_disabled_reason: Optional[str] = None
        if use_rag and not labor_applicable:
            rag_disabled_reason = "Labor-law RAG disabled because Egyptian labor scope is not applicable for this contract."

        # dynamic scopes
        scopes_to_run: List[List[str]] = []
        if labor_applicable:
            scopes_to_run.append(["labor"])
        if is_cb:
            scopes_to_run.append(["cross_border"])

        # ============================================================
        # 2.8) HYBRID: Rule engine (source of truth) + ML assist (severity/prioritization)
        # ============================================================
        rule_hits: List[Dict[str, Any]] = []
        ml_predictions: List[Dict[str, Any]] = []
        ml_used = False
        rag_drove_rules = False  # Rule engine is authoritative when used

        # ---------- 2.9a) Rule engine first (authoritative violations) ----------
        if rule_engine is not None and full_text_rules.strip():
            try:
                flat_scope = []
                for s in scopes_to_run:
                    flat_scope.extend(s)
                engine_hits = rule_engine.check_text(
                    full_text_rules,
                    law_scope=flat_scope or None,
                    contract_type=contract_type,
                    contract_tags=contract_tags,
                )
                if engine_hits:
                    for h in engine_hits:
                        h["chunk_id"] = h.get("chunk_id") or (ocr_chunks[0].get("id") if ocr_chunks else "page_0")
                    rule_hits = engine_hits
            except Exception as e:
                print(f"[RuleEngine][WARN] {e!r}")

        # ---------- 2.9b) ML assist: add severity/priority or fallback when no rule engine hits (only when use_ml=True) ----------
        if use_ml and _HAS_MODEL_ML_PREDICTOR and model_ml_predict_rule_scores_full is not None and model_ml_rule_scores_to_rule_hits is not None:
            try:
                full_text_rules_15k = full_text[:15000] if len(full_text) > 15000 else full_text
                preds_full = model_ml_predict_rule_scores_full(full_text_rules_15k, sort=True, use_law_retrieval=True)
                ml_predictions = preds_full
                ml_used = True
                # If we have rule-engine hits, attach ML scores for severity/prioritization only
                if rule_hits:
                    rid_to_score = {p.get("rule_id"): p.get("score") for p in (preds_full or []) if p.get("rule_id")}
                    for h in rule_hits:
                        h["ml_severity_score"] = rid_to_score.get(h.get("rule_id"))
                else:
                    # No rule engine: use ML hits as fallback (backward compatible)
                    first_chunk_id = ocr_chunks[0].get("id") if ocr_chunks else "page_0"
                    rule_hits = model_ml_rule_scores_to_rule_hits(preds_full, chunk_id=first_chunk_id)
                    for c in ocr_chunks:
                        text_norm = c.get("normalized_text") or ""
                        if len(text_norm.strip()) < 40:
                            continue
                        preds = model_ml_predict_rule_scores_full(text_norm, sort=True, use_law_retrieval=True)
                        chunk_id = c.get("id") or "page_0"
                        hits = model_ml_rule_scores_to_rule_hits(preds, chunk_id=chunk_id)
                        seen_rid_chunk = {(x.get("rule_id"), x.get("chunk_id")) for x in rule_hits}
                        for h in hits:
                            if (h.get("rule_id"), h.get("chunk_id")) not in seen_rid_chunk:
                                rule_hits.append(h)
                                seen_rid_chunk.add((h.get("rule_id"), h.get("chunk_id")))
            except Exception as e:
                print(f"[ML][WARN] Model-ML prediction failed: {e!r}")
                ml_used = False
                ml_predictions = []

        # 3.1) Deduplicate
        seen = set()
        deduped_list: List[Dict[str, Any]] = []
        for h in rule_hits:
            key = (h.get("rule_id"), h.get("chunk_id"), h.get("matched_text"))
            if key in seen:
                continue
            seen.add(key)
            deduped_list.append(h)
        rule_hits = deduped_list

        # 3.5) Summaries
        labor_summary = _build_labor_summary(rule_hits)
        labor_summary = _apply_leave_waiver_to_summary(labor_summary, rule_hits)
        cross_border_summary = _build_cross_border_summary(rule_hits, contract_tags, cb)

        if not labor_applicable:
            labor_summary = {
                "enabled": False,
                "status": "not_applicable",
                "message": "قانون العمل المصري غير منطبق لأن العقد يخضع لقانون/اختصاص أجنبي (مذكور صراحة أو بإشارات قوية).",
                "reason": labor_appx.get("reason"),
                "jurisdiction": labor_appx.get("jurisdiction"),
                "is_cross_border": labor_appx.get("is_cross_border"),
            }

        # ============================================================
        # 4) Query: local matches + RAG (manual query)
        # ============================================================
        rag_results: List[Dict[str, Any]] = []
        contract_local_matches: List[Dict[str, Any]] = []

        if query:
            qnorm = norm_ar(query)
            for c in ocr_chunks:
                if qnorm and qnorm in (c.get("normalized_text") or ""):
                    contract_local_matches.append({"chunk_id": c.get("id"), "text": c.get("text")})

            if use_rag and retriever and labor_applicable:
                boosted_query = build_rag_query(query, full_text, contract_tags)

                rag_results = _rag_search_safe(
                    boosted_query,
                    top_k=8,
                    law_filters={"source": "labor14_2025"},
                    min_score=0.05,
                )

                if not rag_results:
                    rag_results = _rag_search_safe(
                        boosted_query,
                        top_k=8,
                        law_filters={"law": "قانون العمل رقم 14 لسنة 2025"},
                        min_score=0.05,
                    )

                if not rag_results:
                    rag_results = _rag_search_safe(boosted_query, top_k=8, law_filters=None, min_score=0.05)

                rag_results = _filter_to_labor_only(rag_results)[:5]

        # ============================================================
        # 4.2) RAG by violations
        # ============================================================
        rag_by_violation: List[Dict[str, Any]] = []
        rag_global_hits: List[Dict[str, Any]] = []

        def _mk_query_from_violation(h: Dict[str, Any]) -> str:
            rid = (h.get("rule_id") or "").upper()
            desc = h.get("description") or ""
            matched = h.get("matched_text") or ""

            matched_short = (matched or "").strip()
            if len(matched_short) > 220:
                matched_short = matched_short[:220]

            boosts: List[str] = []
            if "WORKING_HOURS" in rid:
                boosts = ["ساعات العمل", "الحد الأقصى", "8 ساعات", "الراحة", "ساعات العمل اليومية", "قانون العمل 14 لسنة 2025"]
            elif "ANNUAL_LEAVE" in rid:
                boosts = ["الإجازة السنوية", "الحد الأدنى", "15 يوم", "21 يوم", "30 يوم", "قانون العمل 14 لسنة 2025"]
            elif "PROBATION" in rid:
                boosts = ["فترة الاختبار", "لا تزيد عن", "3 أشهر", "قانون العمل 14 لسنة 2025"]
            elif "SALARY" in rid:
                boosts = ["الأجر", "طريقة الصرف", "موعد صرف الأجر", "استقطاعات", "تأمينات", "ضرائب", "قانون العمل 14 لسنة 2025"]
            elif "EMPLOYER_INFO" in rid:
                boosts = ["بيانات صاحب العمل", "الاسم", "العنوان", "مقر ممارسة النشاط", "قانون العمل 14 لسنة 2025"]
            elif "EMPLOYEE_INFO" in rid:
                boosts = ["بيانات العامل", "الاسم", "الرقم القومي", "محل الإقامة", "المهنة", "قانون العمل 14 لسنة 2025"]
            elif "PLACEHOLDER" in rid:
                boosts = ["بيانات غير مستكملة", "نقاط", "استكمال البيانات", "بيانات الأطراف", "قانون العمل 14 لسنة 2025"]

            q = " ".join([desc] + boosts + ([matched_short] if matched_short else []))
            return norm_ar(q)

        def _dedupe_hits(hits: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
            seen2 = set()
            out: List[Dict[str, Any]] = []
            for x in hits or []:
                md = x.get("metadata") or {}
                key = (md.get("law"), md.get("article"), md.get("id"), (x.get("text") or "")[:80])
                if key in seen2:
                    continue
                seen2.add(key)
                out.append(x)
            return out

        if use_rag and retriever and labor_applicable:
            important = [h for h in rule_hits if (h.get("severity") in ["error", "high"])][:6]

            for vh in important:
                qv = _mk_query_from_violation(vh)
                qv_boosted = build_rag_query(qv, full_text, contract_tags)

                hits_v = _rag_search_safe(
                    qv_boosted,
                    top_k=6,
                    law_filters={"source": "labor14_2025"},
                    min_score=0.05,
                )

                if not hits_v:
                    hits_v = _rag_search_safe(
                        qv_boosted,
                        top_k=6,
                        law_filters={"law": "قانون العمل رقم 14 لسنة 2025"},
                        min_score=0.05,
                    )

                if not hits_v:
                    hits_v = _rag_search_safe(qv_boosted, top_k=6, law_filters=None, min_score=0.05)

                hits_v = _filter_to_labor_only(_dedupe_hits(hits_v))[:4]

                rag_by_violation.append({
                    "rule_id": vh.get("rule_id"),
                    "severity": vh.get("severity"),
                    "violation_desc": vh.get("description"),
                    "rag_query_used": qv_boosted,
                    "hits": hits_v,
                })
                rag_global_hits.extend(hits_v)

            rag_global_hits = _dedupe_hits(rag_global_hits)
            rag_by_violation = _dedupe_rag_blocks_by_rule(rag_by_violation)

        # ---------- Optional: Local LLM explanation per violation (explanation only, no decision) ----------
        _llm_available = (_llm_explain_violation is not None and _llm_is_available()) or (
            _local_llm_module is not None and getattr(_local_llm_module, "is_available", lambda: False)()
        )
        if use_llm and _llm_available:
            top_k_llm = max(0, int(llm_top_k or 0))
            max_new = max(32, int(llm_max_new_tokens or 200))
            _explain_fn = _llm_explain_violation if _llm_explain_violation is not None else getattr(_local_llm_module, "explain_violation", None)
            for block in (rag_by_violation or [])[:top_k_llm]:
                rule_id = block.get("rule_id") or ""
                violation_desc = block.get("violation_desc") or ""
                hits_v = block.get("hits") or []
                matched = ""
                for h in (rule_hits or []):
                    if h.get("rule_id") == rule_id:
                        matched = (h.get("matched_text") or "")[:800]
                        break
                law_articles = [{"text": x.get("text", ""), "metadata": x.get("metadata") or {}} for x in hits_v]
                try:
                    if _explain_fn:
                        expl = _explain_fn(
                            rule_id=rule_id,
                            description=violation_desc,
                            matched_text=matched or violation_desc,
                            law_articles=law_articles,
                            max_new_tokens=max_new,
                        )
                    else:
                        expl = "[LLM not available]"
                    block["llm_explanation"] = expl
                except Exception as e:
                    block["llm_explanation"] = f"[LLM error: {e!r}]"

        # Dynamic law_scope_used
        law_scope_used: List[str] = []
        if labor_applicable:
            law_scope_used.append("labor")
        if is_cb:
            law_scope_used.append("cross_border")

        # Pipeline audit (plan Step 5): which steps ran
        pipeline_steps: List[str] = ["ocr"]
        if rule_engine is not None and full_text_rules.strip():
            pipeline_steps.append("rule_engine")
        if ml_used:
            pipeline_steps.append("ml_assist")
        if use_rag:
            pipeline_steps.append("rag")
        if use_llm and _llm_available:
            pipeline_steps.append("llm")

        response: Dict[str, Any] = {
            "ocr_chunks": ocr_chunks,
            "clauses": clauses_for_response,
            "rule_hits": rule_hits,
            "labor_summary": labor_summary,
            "cross_border_summary": cross_border_summary,
            "contract_type": contract_type,
            "contract_tags": contract_tags,
            "rag_legal_hits": rag_results,
            "contract_local_matches": contract_local_matches,
            "rag_by_violation": rag_by_violation,
            "rag_global_hits": rag_global_hits,
            "law_scope_used": law_scope_used,
            "cross_border_detector": cb,
            "labor_applicability": labor_appx,
            "rag_disabled_reason": rag_disabled_reason,

            # Non-breaking ML metadata
            "ml_used": ml_used,
            "ml_predictions": ml_predictions,
            "rag_drove_rules": rag_drove_rules,

            # Unified rule+ML: clause-level and full-text violation risk
            "clause_level_unified_risks": clause_level_unified_risks,
            "full_text_unified_risk": full_text_unified_risk,

            # Pipeline audit (structured, auditable)
            "pipeline_steps": pipeline_steps,
        }

        response["cross_border_evidence"] = _rule_evidence(
            rule_hits,
            [
                "CROSSBORDER_GOVERNING_LAW_PRESENT",
                "CROSSBORDER_JURISDICTION_PRESENT",
                "CROSSBORDER_PAYMENT_METHOD_INTERNATIONAL",
                "CROSSBORDER_TRANSFER_FEES_CLARITY",
                "CROSSBORDER_TAX_INSURANCE_CLARITY",
                "CROSSBORDER_REMOTE_TOOLS_SECURITY",
            ],
        )

        # ============================================================
        # Save (optional)
        # ============================================================
        if save:
            if not _HAS_PERSIST or Analysis is None:
                raise HTTPException(status_code=503, detail="Save requested but DB/Auth not configured.")
            if not getattr(current_user, "id", None):
                raise HTTPException(status_code=401, detail="Unauthorized.")

            sha256 = hashlib.sha256(data).hexdigest()
            detected_lang = detect_language(full_text)

            a = Analysis(
                user_id=current_user.id,
                filename=file.filename or "uploaded_contract",
                result_json=json.dumps(response, ensure_ascii=False),
                mime_type=(file.content_type or "application/octet-stream"),
                sha256=sha256,
                page_count=len(ocr_chunks),
                ocr_used=1 if ocr_used_flag else 0,
                detected_lang=detected_lang,
            )
            db.add(a)
            db.commit()
            db.refresh(a)
            response["analysis_id"] = a.id

        return JSONResponse(content=response, media_type="application/json; charset=utf-8")

    except HTTPException:
        raise
    except Exception as e:
        print("ERROR in OCR+Check:", repr(e))
        raise HTTPException(
            status_code=500,
            detail=f"OCR+Check failed: {e.__class__.__name__}: {e}",
        )